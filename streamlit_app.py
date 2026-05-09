import streamlit as st
from supabase import create_client, Client
import hashlib
import sys
import io
import docx
import pandas as pd
import google.generativeai as genai
import os 

# --- Supabase データベース設定 ---

@st.cache_resource 
def init_supabase_client():
    """Supabaseクライアントを初期化して返す"""
    try:
        supabase_url = st.secrets["SUPABASE_URL"]
        supabase_key = st.secrets["SUPABASE_KEY"]
        return create_client(supabase_url, supabase_key)
    except KeyError:
        st.error("Supabase の URL または Key が Streamlit Secrets に設定されていません。")
        st.stop()

def hash_password(password):
    """パスワードをハッシュ化する"""
    return hashlib.sha256(password.encode()).hexdigest()

def add_user(supabase: Client, username, password):
    """一般ユーザーを Supabase に追加する"""
    if username.lower() == 'adminkaho1020':
        return False
    try:
        supabase.table('users').insert({
            'username': username,
            'password_hash': hash_password(password),
            'is_admin': False
        }).execute()
        return True
    except Exception as e:
        st.error(f"不明なエラーが発生しました: {e}")
        return False

def verify_user(supabase: Client, username, password):
    """ユーザーを認証する"""
    try:
        response = supabase.table('users').select('*').eq('username', username).execute()
        if response.data:
            user = response.data[0]
            if user['password_hash'] == hash_password(password):
                return user
        return None
    except Exception as e:
        st.error(f"認証エラー: {e}")
        return None

def get_all_users(supabase: Client):
    """管理者以外の全ユーザーを取得する"""
    try:
        response = supabase.table('users').select('id, username').eq('is_admin', False).order('username').execute()
        return response.data 
    except Exception as e:
        st.error(f"ユーザー取得エラー: {e}")
        return []

def add_message_to_db(supabase: Client, user_id, role, content):
    """チャット履歴を Supabase に追加する"""
    try:
        supabase.table('chat_history').insert({
            'user_id': user_id,
            'role': role,
            'content': content
        }).execute()
    except Exception as e:
        st.error(f"メッセージ保存エラー: {e}")

def get_messages_from_db(supabase: Client, user_id):
    """特定のユーザーのチャット履歴を取得する"""
    try:
        response = supabase.table('chat_history').select('role, content').eq('user_id', user_id).order('timestamp', desc=False).execute()
        return response.data
    except Exception as e:
        st.error(f"履歴取得エラー: {e}")
        return []

# ★★★ 新規追加：過去データの取得（プレースホルダー） ★★★
def get_past_learning_record(supabase: Client, user_id):
    """
    過去の学習記録を取得する。
    ※現在は将来のためのプレースホルダーとして「データなし」を返します。
    将来的に、DBから過去の要約データを取得するロジックをここに実装します。
    """
    return {
        "challenge": "データなし",
        "achievement": "データなし"
    }

# --- 管理者パネル ---
def admin_panel(supabase: Client): 
    st.sidebar.title("管理者パネル")
    st.sidebar.write("---")
    
    if st.session_state.get('impersonating', False):
        if st.sidebar.button("管理者ビューに戻る"):
            st.session_state['user_id'] = st.session_state['admin_id']
            st.session_state['username'] = st.session_state['admin_username']
            st.session_state['is_admin'] = True
            st.session_state['impersonating'] = False
            if 'viewing_messages' in st.session_state:
                del st.session_state['viewing_messages']
            st.rerun()
        st.sidebar.write("---")

    st.sidebar.subheader("ユーザー一覧")
    users = get_all_users(supabase) 
    if not users:
        st.sidebar.info("まだ一般ユーザーは登録されていません。")
        return

    for user in users:
        with st.sidebar.expander(f"ユーザー: {user['username']}"):
            if st.button("履歴を閲覧", key=f"view_{user['id']}"):
                messages = get_messages_from_db(supabase, user['id']) 
                st.session_state['viewing_messages'] = messages
                st.session_state['viewing_username'] = user['username']
                if 'impersonating' in st.session_state:
                    st.session_state['impersonating'] = False

            if st.button("このユーザーとしてログイン", key=f"login_as_{user['id']}"):
                st.session_state['impersonating'] = True
                st.session_state['admin_id'] = st.session_state['user_id']
                st.session_state['admin_username'] = st.session_state['username']
                st.session_state['user_id'] = user['id']
                st.session_state['username'] = user['username']
                st.session_state['is_admin'] = False
                st.session_state.messages = get_messages_from_db(supabase, user['id']) 
                if 'viewing_messages' in st.session_state:
                    del st.session_state['viewing_messages']
                st.rerun()


# --- メインアプリケーション ---
def main():
    supabase = init_supabase_client()

    # セッション状態の初期化
    if 'logged_in' not in st.session_state:
        st.session_state.logged_in = False
        st.session_state.username = ""
        st.session_state.user_id = None
        st.session_state.is_admin = False

    # --- ログイン/新規登録UI (サイドバー) ---
    if not st.session_state.logged_in:
        st.sidebar.title("ユーザー認証")
        choice = st.sidebar.selectbox("メニュー", ["ログイン", "新規登録"])

        if choice == "ログイン":
            with st.sidebar.form("login_form"):
                username = st.text_input("ユーザー名")
                password = st.text_input("パスワード", type="password")
                submitted = st.form_submit_button("ログイン")
                if submitted:
                    user = verify_user(supabase, username, password) 
                    if user:
                        st.session_state.logged_in = True
                        st.session_state.username = user['username']
                        st.session_state.user_id = user['id']
                        st.session_state.is_admin = user['is_admin']
                        st.rerun()
                    else:
                        st.sidebar.error("ユーザー名またはパスワードが間違っています。")

        elif choice == "新規登録":
            with st.sidebar.form("signup_form"):
                new_username = st.text_input("ユーザー名")
                if new_username.lower() == 'adminkaho1020':
                    st.warning("このユーザー名は使用できません。")
                new_password = st.text_input("パスワード", type="password")
                submitted = st.form_submit_button("登録")
                if submitted and new_username.lower() != 'adminkaho1020':
                    if add_user(supabase, new_username, new_password): 
                        st.sidebar.success("登録が完了しました。ログインしてください。")
                    else:
                        st.sidebar.error("このユーザー名は既に使用されているか、登録に失敗しました。")
    else: 
        st.sidebar.success(f"{st.session_state.username} としてログイン中")
        if st.sidebar.button("ログアウト"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()

    # --- ログインしている場合のみアプリ本体を表示 ---
    if st.session_state.logged_in:
        if st.session_state.is_admin and not st.session_state.get('impersonating', False):
            admin_panel(supabase) 
            st.title("管理者ダッシュボード")
            st.info("サイドバーからユーザーを選択し、操作を行ってください。")

            if 'viewing_messages' in st.session_state:
                st.header(f"ユーザー「{st.session_state['viewing_username']}」の学習履歴")
                messages_to_display = st.session_state['viewing_messages']
                if not messages_to_display:
                    st.write("このユーザーのチャット履歴はまだありません。")
                else:
                    for message in messages_to_display:
                        with st.chat_message(message["role"]):
                            st.markdown(message["content"])
        
        else:
            if st.session_state.get('impersonating', False):
                st.info(f"現在、管理者として「{st.session_state.username}」でログインしています。")
                if st.sidebar.button("管理者ビューに戻る"):
                    st.session_state.user_id = st.session_state.admin_id
                    st.session_state.username = st.session_state.admin_username
                    st.session_state.is_admin = True
                    st.session_state.impersonating = False
                    if 'viewing_messages' in st.session_state:
                        del st.session_state['viewing_messages']
                    st.rerun()
            
            # --- チャットアプリ本体 ---
            st.title("💬 チャットボットと学びを振り返ろう！")
            st.write("記入済みの学習日記フォーマットをDOCS形式でアップロードすると、その内容に関する対話ができます！")

            try:
                gemini_api_key = st.secrets["google_api_key"]
                genai.configure(api_key=gemini_api_key)
                
                # ★★★ ベースシステムプロンプトの定義（変数を埋め込める形） ★★★
                base_system_prompt = """
あなたはユーザーがアップロードしたファイル内の「学習目標」に精通した優秀な指導教員であり、独学する成人学習者の自己成長を支援する親しみやすいコーチングチャットボットです。

### 0. リアルタイム・コンテキスト（動的注入データ）
以下の情報はシステムによって自動更新されます。これらに基づいて回答を調整してください。

<past_learning_record>
・前回の主要な課題: {past_challenge}
・前回達成したこと: {past_achievement}
</past_learning_record>

<user_status>
・直近の入力の長さ: {input_length_status}
・疲労フラグ: {fatigue_flag}
</user_status>

---

### 1. 最重要ルール：フェーズ制と疲労への配慮
対話履歴の往復回数を確認し、以下のフェーズを厳密に守ってください。ただし、<user_status>の「疲労フラグ」が ON の場合は、即座にフェーズ3（終了）へ誘導してください。

* **フェーズ1（1〜3往復目）：徹底的な内省（Step 1）**
    - 目的：安易な解決策を出さず、ユーザーの思考を深く掘り下げる。
    - 義務：3往復目の最後に「現在ステップ1/3が終了です。次は成長の振り返りですが、続けても大丈夫ですか？」と進捗を確認すること。

* **フェーズ2（4〜6往復目）：視点の転換と過去比較（Step 2）**
    - 目的：過去の自分（<past_learning_record>）と比較し、成長を実感させる。
    - 義務：必ず「前回の課題であった〇〇が、今回は△△になっていますね」と言及すること。

* **フェーズ3（7往復目以降、または終了希望時）：行動への橋渡し（Step 3）**
    - 目的：次回の具体的な一歩（Volition）を決め、ポジティブに締めくくる。

---

### 2. 対話の進行プロセス（Step by Step）

#### 【ステップ1：深掘り】
ユーザーの回答に対し、「なぜ？」「具体的には？」と**最低2回以上**質問を重ねてください。
※疲労フラグがON、または入力が極端に短い場合は、深掘りを中止して労りの言葉をかけてください。

#### 【ステップ2：過去比較による自信の醸成】
`<past_learning_record>` を参照し、以下の構成で話してください。
1. **過去の引用:** 「前回は[課題]と仰っていましたが、」
2. **成長の承認:** 「今回は[今回の気づき]ができていますね！素晴らしい進歩です。」
※データが「データなし」の場合は、本日の対話の冒頭の発言と比較してください。

#### 【ステップ3：意思確認とクロージング】
次回の学習に向けた具体的な行動計画をユーザーに宣言させ、「次回も楽しみにしています！」とポジティブに終了します。

---

### 3. 禁止事項・スタンス
* 直接的なアドバイスや正解の提示は行わない。
* 専門的な質問には「一緒に調べましょう」または「検索を促す」に留める。
* 常にARCS-Vモデルを意識し、自信（C）と意志（V）を高める声掛けを徹底する。
"""
            except Exception as e:
                st.error(f"APIキーの設定でエラーが発生しました: {e}")
                st.stop()
            
            uploaded_file = st.file_uploader("ドキュメントをアップロードしてください", type=['txt', 'docx'])

            if "messages" not in st.session_state:
                st.session_state.messages = get_messages_from_db(supabase, st.session_state.user_id) 
            if "document_content" not in st.session_state:
                st.session_state.document_content = None

            # --- ファイルアップロード時の処理（初回起動） ---
            if uploaded_file is not None and st.session_state.document_content is None:
                 try:
                    if uploaded_file.type == 'text/plain':
                        document_content = uploaded_file.getvalue().decode('utf-8')
                    elif uploaded_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                        document = docx.Document(uploaded_file)
                        paragraphs = [p.text for p in document.paragraphs]
                        document_content = "\n".join(paragraphs)
                    
                    st.session_state.document_content = document_content
                    st.success("ドキュメントが正常にアップロードされました。")
                    st.info("これで、ドキュメントの内容について質問できます。")
                    
                    initial_prompt = f"""
あなたは今、システムプロンプト（役割定義）に従い、指導教員/コーチとして振る舞っています。
学習者（ユーザー）が、以下の学習日記（ドキュメント）をアップロードしました。
このドキュメントの内容を解釈し、システムプロンプトの「ステップ1」に従って、最初の応答（Botラリー1）を生成してください。
ワンパターンな質問ではなく、日記の内容に具体的に言及し、回答しやすい具体的な問いかけを心がけてください。

---
学習日記（ドキュメント）:
{document_content}
---

あなたの最初の応答を開始してください：
"""
                    
                    with st.spinner("思考中です..."):
                        # ★★★ 初回のモデル初期化（デフォルト値を代入） ★★★
                        p_data = get_past_learning_record(supabase, st.session_state.user_id)
                        current_system_prompt = base_system_prompt.format(
                            past_challenge=p_data['challenge'],
                            past_achievement=p_data['achievement'],
                            input_length_status="適切", # 初回は適切とする
                            fatigue_flag="OFF"          # 初回は疲労なし
                        )
                        model = genai.GenerativeModel('gemini-2.5-flash', system_instruction=current_system_prompt)
                        response = model.generate_content(initial_prompt)
                    
                    assistant_message = response.text
                    st.session_state.messages.append({"role": "assistant", "content": assistant_message})
                    add_message_to_db(supabase, st.session_state['user_id'], "assistant", assistant_message)
                    st.rerun()
                 except Exception as e:
                    st.error(f"ファイルの読み込み中にエラーが発生しました: {e}")

            # --- 既存のメッセージを表示 ---
            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

            # --- ユーザーの入力があった場合の処理 ---
            if prompt := st.chat_input("ドキュメントについて質問してください"):
                st.session_state.messages.append({"role": "user", "content": prompt})
                add_message_to_db(supabase, st.session_state.user_id, "user", prompt) 
                with st.chat_message("user"):
                    st.markdown(prompt)

                try:
                    # ★★★ 疲労検知＆文字数判定ロジック ★★★
                    input_len = len(prompt)
                    input_length_status = "短い" if input_len < 10 else "適切"
                    
                    # session_state.messages にはユーザーとAIの会話が蓄積されているため、
                    # length が 6 以上であれば（最低でも3往復目に入っている）
                    session_turn_count = len(st.session_state.messages)
                    fatigue_flag = "ON" if (input_len < 10 and session_turn_count >= 6) else "OFF"

                    # ★★★ 過去データの取得とプロンプトの動的生成 ★★★
                    p_data = get_past_learning_record(supabase, st.session_state.user_id)
                    dynamic_system_prompt = base_system_prompt.format(
                        past_challenge=p_data['challenge'],
                        past_achievement=p_data['achievement'],
                        input_length_status=input_length_status,
                        fatigue_flag=fatigue_flag
                    )

                    # ★★★ モデルを最新のシステムプロンプトで再初期化 ★★★
                    model = genai.GenerativeModel('gemini-2.5-flash', system_instruction=dynamic_system_prompt)

                    # 履歴構築
                    history = []
                    document_context = f"参考：ユーザーの学習日記（ドキュメント）:\n{st.session_state.get('document_content', 'ドキュメントなし')}"
                    history.append({'role': 'user', 'parts': [document_context]})
                    history.append({'role': 'model', 'parts': ["（承知しました。学習日記を再度参照します。）"]})

                    for msg in st.session_state.messages:
                        role = "user" if msg["role"] == "user" else "model"
                        history.append({'role': role, 'parts': [msg["content"]]})
                    
                    # 生成
                    response_stream = model.generate_content(history, stream=True)

                    full_response = ""
                    with st.chat_message("assistant"):
                        message_placeholder = st.empty()
                        for chunk in response_stream:
                            if chunk.parts:
                                text_part = chunk.parts[0].text
                                full_response += text_part
                                message_placeholder.markdown(full_response + "▌")
                        message_placeholder.markdown(full_response)
                    
                    st.session_state.messages.append({"role": "assistant", "content": full_response})
                    add_message_to_db(supabase, st.session_state.user_id, "assistant", full_response) 

                except Exception as e:
                    st.error("エラーが発生しました。詳細はコンソールを確認してください。")
                    print(f"エラーの詳細: {e}", file=sys.stderr)
                    error_message = "申し訳ありません、応答の生成中にエラーが発生しました。"
                    st.session_state.messages.append({"role": "assistant", "content": error_message})
                    add_message_to_db(supabase, st.session_state.user_id, "assistant", error_message) 
            
            # --- エクスポート機能 ---
            st.sidebar.header("エクスポート")
            doc = docx.Document()
            doc.add_heading(f'{st.session_state["username"]}さんの振り返り', 0)
            for message in st.session_state.messages:
                role_jp = "ユーザー" if message["role"] == "user" else "チャットボット"
                doc.add_paragraph(f"{role_jp}: {message['content']}")
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            st.sidebar.download_button(
                label="振り返りをWord形式でダウンロード",
                data=doc_io,
                file_name=f"{st.session_state['username']}_振り返り.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            if st.session_state.messages:
                df = pd.DataFrame(st.session_state.messages)
                csv = df.to_csv(index=False).encode('utf-8')
                st.sidebar.download_button(
                    label="対話履歴をCSV形式でダウンロード",
                    data=csv,
                    file_name=f"{st.session_state['username']}_対話履歴.csv",
                    mime="text/csv",
                )
    else:
        st.info("チャットボットを利用するには、サイドバーからログインまたは新規登録をしてください。")

if __name__ == '__main__':
    main()
